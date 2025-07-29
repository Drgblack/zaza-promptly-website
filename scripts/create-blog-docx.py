#!/usr/bin/env python3
"""
Convert blog markdown to teacher-friendly Word document
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_blog_docx(md_file_path, output_path, title_override=None):
    """Convert blog markdown to professional Word document"""
    
    # Read markdown content
    with open(md_file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Extract frontmatter
    frontmatter_match = re.match(r'^---\n(.*?)\n---\n(.*)', content, re.DOTALL)
    if frontmatter_match:
        frontmatter, main_content = frontmatter_match.groups()
        
        # Extract title from frontmatter
        title_match = re.search(r'title:\s*["\']?(.*?)["\']?\s*$', frontmatter, re.MULTILINE)
        title = title_match.group(1) if title_match else "Zaza Promptly Blog Post"
        
        # Extract description
        desc_match = re.search(r'description:\s*["\']?(.*?)["\']?\s*$', frontmatter, re.MULTILINE)
        description = desc_match.group(1) if desc_match else ""
        
        # Extract reading time
        time_match = re.search(r'readingTime:\s*["\']?(.*?)["\']?\s*$', frontmatter, re.MULTILINE)
        reading_time = time_match.group(1) if time_match else "5 min read"
    else:
        main_content = content
        title = title_override or "Zaza Promptly Blog Post"
        description = ""
        reading_time = "5 min read"
    
    # Create new document
    doc = Document()
    
    # Set up styles
    setup_blog_styles(doc)
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
    
    # Add header with Zaza branding
    header_p = doc.add_paragraph()
    header_p.style = 'Blog Header'
    header_run = header_p.add_run("ZAZA PROMPTLY EDUCATOR COMMUNITY")
    header_run.font.size = Pt(10)
    header_run.font.name = 'Calibri'
    header_run.bold = True
    
    # Add title
    title_p = doc.add_paragraph(title)
    title_p.style = 'Blog Title'
    
    # Add description if exists
    if description:
        desc_p = doc.add_paragraph(description)
        desc_p.style = 'Blog Subtitle'
    
    # Add reading time
    meta_p = doc.add_paragraph(f"📖 {reading_time} • For Educators, By Educators")
    meta_p.style = 'Blog Meta'
    
    # Add separator
    doc.add_paragraph("─" * 50).style = 'Blog Separator'
    
    # Process main content
    lines = main_content.split('\n')
    
    for line in lines:
        line = line.strip()
        
        if not line:
            continue
            
        # Handle headers
        if line.startswith('# '):
            continue  # Skip main title (already added)
        elif line.startswith('## '):
            p = doc.add_paragraph(line[3:])
            p.style = 'Blog H2'
        elif line.startswith('### '):
            p = doc.add_paragraph(line[4:])
            p.style = 'Blog H3'
        
        # Handle emphasis
        elif line.startswith('*') and line.endswith('*') and len(line) > 2:
            p = doc.add_paragraph()
            p.style = 'Blog Emphasis'
            run = p.add_run(line[1:-1])
            run.italic = True
        
        # Handle bullet points
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
            p.style = 'Blog List'
            
        # Handle checkboxes
        elif line.startswith('- ✅'):
            p = doc.add_paragraph(line[2:])
            p.style = 'Blog List'
            
        # Handle bold emphasis  
        elif '**' in line:
            p = doc.add_paragraph()
            p.style = 'Blog Body'
            
            # Split by ** and alternate between normal and bold
            parts = line.split('**')
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 == 1:  # Odd indices are bold
                    run.bold = True
        
        # Regular paragraph
        elif line and not line.startswith('#'):
            p = doc.add_paragraph(line)
            p.style = 'Blog Body'
    
    # Add closing section
    doc.add_paragraph("─" * 50).style = 'Blog Separator'
    
    # Add footer
    footer_p = doc.add_paragraph()
    footer_p.style = 'Blog Footer'
    footer_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    footer_run1 = footer_p.add_run("Thank you for being part of the Zaza Promptly educator community.\n")
    footer_run1.font.size = Pt(10)
    
    footer_run2 = footer_p.add_run("Download more resources at ZazaPromptly.com/free-resources")
    footer_run2.font.size = Pt(9)
    footer_run2.italic = True
    
    # Save the document
    doc.save(output_path)
    print(f"Created: {output_path}")

def setup_blog_styles(doc):
    """Set up professional blog document styles"""
    
    # Blog Header style
    header_style = doc.styles.add_style('Blog Header', WD_STYLE_TYPE.PARAGRAPH)
    header_font = header_style.font
    header_font.name = 'Calibri'
    header_font.size = Pt(10)
    header_font.color.rgb = None
    header_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_style.paragraph_format.space_after = Pt(12)
    
    # Blog Title style
    title_style = doc.styles.add_style('Blog Title', WD_STYLE_TYPE.PARAGRAPH)
    title_font = title_style.font
    title_font.name = 'Calibri'
    title_font.size = Pt(22)
    title_font.bold = True
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(12)
    title_style.paragraph_format.space_before = Pt(6)
    
    # Blog Subtitle style
    subtitle_style = doc.styles.add_style('Blog Subtitle', WD_STYLE_TYPE.PARAGRAPH)
    subtitle_font = subtitle_style.font
    subtitle_font.name = 'Calibri'
    subtitle_font.size = Pt(14)
    subtitle_font.italic = True
    subtitle_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_style.paragraph_format.space_after = Pt(8)
    
    # Blog Meta style
    meta_style = doc.styles.add_style('Blog Meta', WD_STYLE_TYPE.PARAGRAPH)
    meta_font = meta_style.font
    meta_font.name = 'Calibri'
    meta_font.size = Pt(10)
    meta_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_style.paragraph_format.space_after = Pt(18)
    
    # Blog Separator style
    sep_style = doc.styles.add_style('Blog Separator', WD_STYLE_TYPE.PARAGRAPH)
    sep_font = sep_style.font
    sep_font.name = 'Calibri'
    sep_font.size = Pt(12)
    sep_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep_style.paragraph_format.space_before = Pt(12)
    sep_style.paragraph_format.space_after = Pt(12)
    
    # Blog H2 style
    h2_style = doc.styles.add_style('Blog H2', WD_STYLE_TYPE.PARAGRAPH)
    h2_font = h2_style.font
    h2_font.name = 'Calibri'
    h2_font.size = Pt(16)
    h2_font.bold = True
    h2_style.paragraph_format.space_before = Pt(16)
    h2_style.paragraph_format.space_after = Pt(8)
    
    # Blog H3 style
    h3_style = doc.styles.add_style('Blog H3', WD_STYLE_TYPE.PARAGRAPH)
    h3_font = h3_style.font
    h3_font.name = 'Calibri'
    h3_font.size = Pt(14)
    h3_font.bold = True
    h3_style.paragraph_format.space_before = Pt(12)
    h3_style.paragraph_format.space_after = Pt(6)
    
    # Blog Body style
    body_style = doc.styles.add_style('Blog Body', WD_STYLE_TYPE.PARAGRAPH)
    body_font = body_style.font
    body_font.name = 'Calibri'
    body_font.size = Pt(11)
    body_style.paragraph_format.space_after = Pt(6)
    body_style.paragraph_format.line_spacing = 1.15
    
    # Blog Emphasis style
    emphasis_style = doc.styles.add_style('Blog Emphasis', WD_STYLE_TYPE.PARAGRAPH)
    emphasis_font = emphasis_style.font
    emphasis_font.name = 'Calibri'
    emphasis_font.size = Pt(11)
    emphasis_font.italic = True
    emphasis_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    emphasis_style.paragraph_format.space_before = Pt(8)
    emphasis_style.paragraph_format.space_after = Pt(8)
    
    # Blog List style
    list_style = doc.styles.add_style('Blog List', WD_STYLE_TYPE.PARAGRAPH)
    list_font = list_style.font
    list_font.name = 'Calibri'
    list_font.size = Pt(11)
    list_style.paragraph_format.left_indent = Inches(0.25)
    list_style.paragraph_format.space_after = Pt(3)
    
    # Blog Footer style
    footer_style = doc.styles.add_style('Blog Footer', WD_STYLE_TYPE.PARAGRAPH)
    footer_font = footer_style.font
    footer_font.name = 'Calibri'
    footer_font.size = Pt(10)
    footer_style.paragraph_format.space_before = Pt(18)

def main():
    """Convert the new blog post to Word format"""
    
    blog_dir = Path("/workspaces/zaza-websites-claudecode/content/blog")
    output_dir = Path("/workspaces/zaza-websites-claudecode/public/downloads/blogs")
    
    md_file = blog_dir / "when-grading-feels-overwhelming.md"
    docx_file = output_dir / "When_Grading_Feels_Overwhelming.docx"
    
    if md_file.exists():
        create_blog_docx(md_file, docx_file)
    else:
        print(f"Markdown file not found: {md_file}")

if __name__ == "__main__":
    main()