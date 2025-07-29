#!/usr/bin/env python3
"""
Convert teacher resource markdown files to professional Word documents
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn

def add_hyperlink(paragraph, text, url):
    """Add hyperlink to paragraph"""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Add color and underline for hyperlink
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    new_run.text = text
    
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def setup_document_styles(doc):
    """Set up professional document styles"""
    
    # Title style
    title_style = doc.styles.add_style('Custom Title', WD_STYLE_TYPE.PARAGRAPH)
    title_font = title_style.font
    title_font.name = 'Calibri'
    title_font.size = Pt(24)
    title_font.bold = True
    title_font.color.rgb = None  # Default color
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(12)
    
    # Subtitle style
    subtitle_style = doc.styles.add_style('Custom Subtitle', WD_STYLE_TYPE.PARAGRAPH)
    subtitle_font = subtitle_style.font
    subtitle_font.name = 'Calibri'
    subtitle_font.size = Pt(16)
    subtitle_font.bold = True
    subtitle_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_style.paragraph_format.space_after = Pt(18)
    
    # Header 1 style
    h1_style = doc.styles.add_style('Custom H1', WD_STYLE_TYPE.PARAGRAPH)
    h1_font = h1_style.font
    h1_font.name = 'Calibri'
    h1_font.size = Pt(18)
    h1_font.bold = True
    h1_style.paragraph_format.space_before = Pt(18)
    h1_style.paragraph_format.space_after = Pt(12)
    
    # Header 2 style
    h2_style = doc.styles.add_style('Custom H2', WD_STYLE_TYPE.PARAGRAPH)
    h2_font = h2_style.font
    h2_font.name = 'Calibri'
    h2_font.size = Pt(14)
    h2_font.bold = True
    h2_style.paragraph_format.space_before = Pt(12)
    h2_style.paragraph_format.space_after = Pt(6)
    
    # Header 3 style
    h3_style = doc.styles.add_style('Custom H3', WD_STYLE_TYPE.PARAGRAPH)
    h3_font = h3_style.font
    h3_font.name = 'Calibri'
    h3_font.size = Pt(12)
    h3_font.bold = True
    h3_style.paragraph_format.space_before = Pt(6)
    h3_style.paragraph_format.space_after = Pt(6)
    
    # Body text style
    body_style = doc.styles.add_style('Custom Body', WD_STYLE_TYPE.PARAGRAPH)
    body_font = body_style.font
    body_font.name = 'Calibri'
    body_font.size = Pt(11)
    body_style.paragraph_format.space_after = Pt(6)
    body_style.paragraph_format.line_spacing = 1.15

def convert_markdown_to_docx(md_file_path, output_path):
    """Convert markdown file to Word document with professional formatting"""
    
    # Read markdown content
    with open(md_file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create new document
    doc = Document()
    
    # Set up styles
    setup_document_styles(doc)
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Split content into lines
    lines = content.split('\n')
    
    # Track if we're in a code block
    in_code_block = False
    code_content = []
    
    # Process each line
    for line in lines:
        line = line.strip()
        
        # Skip empty lines unless in code block
        if not line and not in_code_block:
            continue
            
        # Handle code blocks
        if line.startswith('```'):
            if in_code_block:
                # End of code block - add as formatted text
                if code_content:
                    p = doc.add_paragraph()
                    p.style = 'Custom Body'
                    p.paragraph_format.left_indent = Inches(0.5)
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(6)
                    
                    # Add code content with monospace font
                    run = p.add_run('\n'.join(code_content))
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                    
                code_content = []
                in_code_block = False
            else:
                in_code_block = True
            continue
        
        # If in code block, collect content
        if in_code_block:
            code_content.append(line)
            continue
        
        # Handle headers
        if line.startswith('# '):
            p = doc.add_paragraph(line[2:])
            p.style = 'Custom Title'
        elif line.startswith('## '):
            p = doc.add_paragraph(line[3:])
            p.style = 'Custom Subtitle'
        elif line.startswith('### '):
            p = doc.add_paragraph(line[4:])
            p.style = 'Custom H1'
        elif line.startswith('#### '):
            p = doc.add_paragraph(line[5:])
            p.style = 'Custom H2'
        elif line.startswith('##### '):
            p = doc.add_paragraph(line[6:])
            p.style = 'Custom H3'
        
        # Handle bullet points
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25)
            
        # Handle numbered lists
        elif re.match(r'^\d+\. ', line):
            content_text = re.sub(r'^\d+\. ', '', line)
            p = doc.add_paragraph(content_text, style='List Number')
            p.paragraph_format.left_indent = Inches(0.25)
            
        # Handle checkboxes
        elif line.startswith('□ [ ] ') or line.startswith('- [ ] '):
            checkbox_text = line.replace('□ [ ] ', '☐ ').replace('- [ ] ', '☐ ')
            p = doc.add_paragraph(checkbox_text)
            p.style = 'Custom Body'
            p.paragraph_format.left_indent = Inches(0.25)
            
        # Handle emphasis and formatting
        elif line:
            p = doc.add_paragraph()
            p.style = 'Custom Body'
            
            # Process bold, italic, and other formatting
            text = line
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Remove bold markers, keep text
            text = re.sub(r'\*(.*?)\*', r'\1', text)      # Remove italic markers, keep text
            
            # Add the text
            run = p.add_run(text)
            
            # Apply bold if originally had ** markers
            if '**' in line:
                run.bold = True
    
    # Add footer with generation info
    footer_p = doc.add_paragraph()
    footer_p.style = 'Custom Body'
    footer_p.paragraph_format.space_before = Pt(24)
    footer_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_p.add_run("Generated by Zaza Promptly • ZazaPromptly.com")
    footer_run.font.size = Pt(9)
    footer_run.italic = True
    
    # Save the document
    doc.save(output_path)
    print(f"Created: {output_path}")

def main():
    """Convert all markdown files to docx"""
    
    downloads_dir = Path("/workspaces/zaza-websites-claudecode/public/downloads")
    
    # Files to convert with new descriptive names
    files_to_convert = {
        "ai-prompt-templates.md": "AI_Prompt_Templates_for_Teachers.docx",
        "assessment-templates.md": "Assessment_Rubrics_and_Templates.docx", 
        "classroom-management-guide.md": "Classroom_Management_Guide.docx",
        "lesson-planning-templates.md": "Lesson_Planning_Templates.docx",
        "time-management-guide.md": "Teacher_Time_Management_Guide.docx",
        "weekly-newsletter-signup.md": "Weekly_Teacher_Newsletter_Info.docx"
    }
    
    for md_file, docx_file in files_to_convert.items():
        md_path = downloads_dir / md_file
        docx_path = downloads_dir / docx_file
        
        if md_path.exists():
            print(f"Converting {md_file} to {docx_file}...")
            convert_markdown_to_docx(md_path, docx_path)
        else:
            print(f"Warning: {md_file} not found")
    
    print("\nConversion complete! All .docx files created in /public/downloads/")

if __name__ == "__main__":
    main()